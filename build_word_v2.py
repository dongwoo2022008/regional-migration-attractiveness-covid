"""
Convert manuscript_final_v2.md to Word with figures embedded.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

# ── path mapping for figures ──────────────────────────────────────────────────
FIGURE_PATHS = {
    'figure_table/fig1_research_framework.png':
        '/home/ubuntu/regional-migration-attractiveness-covid/results/figures/fig1_research_framework.png',
    'figure_table/fig4_15_shap_summary.png':
        '/home/ubuntu/upload/figure_table/fig4_15_shap_summary.png',
    'figure_table/t2_invertedU_density.png':
        '/home/ubuntu/regional-migration-attractiveness-covid/results/figures/t2_invertedU_density.png',
    'figure_table/fig6_cluster_characteristics.png':
        '/home/ubuntu/regional-migration-attractiveness-covid/results/figures/fig6_cluster_characteristics.png',
    'figure_table/fig7_stage_shap.png':
        '/home/ubuntu/regional-migration-attractiveness-covid/results/figures/fig7_stage_shap.png',
    'figure_table/t1_typology_pca.png':
        '/home/ubuntu/regional-migration-attractiveness-covid/results/figures/t1_typology_pca.png',
}

with open('manuscript_final_v2.md', 'r', encoding='utf-8') as f:
    text = f.read()

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── default style ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

lines = text.split('\n')
in_table = False
table_data = []

def flush_table(doc, table_data):
    if not table_data:
        return
    ncols = len(table_data[0])
    table = doc.add_table(rows=len(table_data), cols=ncols)
    table.style = 'Table Grid'
    for i, row in enumerate(table_data):
        for j, cell_text in enumerate(row):
            cell_text = cell_text.replace('**', '')
            tc = table.cell(i, j)
            tc.text = cell_text
            if i == 0:
                for run in tc.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph('')

i = 0
while i < len(lines):
    raw = lines[i]
    line = raw.strip()

    # empty line
    if not line:
        if in_table:
            flush_table(doc, table_data)
            in_table = False
            table_data = []
        else:
            doc.add_paragraph('')
        i += 1
        continue

    # table row
    if line.startswith('|'):
        if '---' not in line:
            in_table = True
            row = [c.strip() for c in line.split('|')[1:-1]]
            table_data.append(row)
        i += 1
        continue
    else:
        if in_table:
            flush_table(doc, table_data)
            in_table = False
            table_data = []

    # headings
    if line.startswith('# '):
        p = doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        p = doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        p = doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        p = doc.add_heading(line[5:], level=4)

    # image
    elif line.startswith('!['):
        match = re.search(r'!\[.*?\]\((.*?)\)', line)
        if match:
            img_key = match.group(1)
            img_path = FIGURE_PATHS.get(img_key, '')
            if img_path and os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5.8))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph(f'[Figure not found: {img_key}]')
        # next line is caption
        if i + 1 < len(lines) and lines[i+1].strip().startswith('*'):
            i += 1
            caption_text = lines[i].strip().strip('*')
            cp = doc.add_paragraph(caption_text)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cp.runs:
                run.italic = True
                run.font.size = Pt(10)

    # bullet
    elif line.startswith('*   ') or line.startswith('* '):
        text_content = line.lstrip('* ')
        p = doc.add_paragraph(style='List Bullet')
        parts = re.split(r'(\*\*.*?\*\*)', text_content)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

    # blockquote
    elif line.startswith('> '):
        p = doc.add_paragraph(line[2:])
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.italic = True

    # horizontal rule
    elif line.startswith('---'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    # normal paragraph
    else:
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

    i += 1

# flush any remaining table
if in_table:
    flush_table(doc, table_data)

doc.save('Manuscript_Final_v2_PSP.docx')
print('Saved: Manuscript_Final_v2_PSP.docx')
